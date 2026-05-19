"""Document text extraction service — PDF, DOCX, TXT, and images (OCR)."""

from dataclasses import dataclass, field
from pathlib import Path

import fitz  # PyMuPDF
import structlog

logger = structlog.get_logger()


def _sanitize_text(text: str) -> str:
    """Remove null bytes and other characters PostgreSQL rejects."""
    return text.replace("\x00", "")


@dataclass
class PageContent:
    """Represents extracted text from a single page."""

    page_number: int
    text: str


@dataclass
class ExtractedDocument:
    """Result of text extraction from a document."""

    pages: list[PageContent] = field(default_factory=list)
    total_pages: int = 0
    title: str = ""

    @property
    def full_text(self) -> str:
        return "\n\n".join(page.text for page in self.pages if page.text.strip())


class TextExtractor:
    """Extracts text from various document formats including OCR for images."""

    async def extract(self, file_path: str, file_type: str) -> ExtractedDocument:
        """Route extraction based on file type."""
        match file_type.lower():
            case "pdf":
                return self._extract_pdf(file_path)
            case "txt" | "md":
                return self._extract_text(file_path)
            case "docx":
                return self._extract_docx(file_path)
            case "png" | "jpg" | "jpeg":
                return self._extract_image_ocr(file_path)
            case _:
                raise ValueError(f"Unsupported file type: {file_type}")

    def _extract_pdf(self, file_path: str) -> ExtractedDocument:
        """Extract text from PDF using PyMuPDF.

        Falls back to OCR for image-based PDFs with no selectable text.
        """
        doc = fitz.open(file_path)
        total_pages = len(doc)
        pages = []

        for page_num in range(total_pages):
            page = doc[page_num]
            text = _sanitize_text(page.get_text("text"))

            if text.strip():
                pages.append(PageContent(page_number=page_num + 1, text=text.strip()))
            else:
                # Try OCR on this page (image-based PDF)
                ocr_text = self._ocr_pdf_page(page)
                if ocr_text.strip():
                    pages.append(PageContent(page_number=page_num + 1, text=ocr_text.strip()))

        title = Path(file_path).stem

        # Try to get title from PDF metadata
        metadata = doc.metadata
        if metadata and metadata.get("title"):
            title = metadata["title"]

        doc.close()

        logger.info(
            "pdf_extracted",
            file_path=file_path,
            total_pages=total_pages,
            pages_with_text=len(pages),
            total_chars=sum(len(p.text) for p in pages),
        )

        return ExtractedDocument(pages=pages, total_pages=total_pages, title=title)

    def _extract_text(self, file_path: str) -> ExtractedDocument:
        """Extract text from plain text files."""
        path = Path(file_path)
        text = _sanitize_text(path.read_text(encoding="utf-8"))

        return ExtractedDocument(
            pages=[PageContent(page_number=1, text=text)],
            total_pages=1,
            title=path.stem,
        )

    def _extract_docx(self, file_path: str) -> ExtractedDocument:
        """Extract text from DOCX files with structure awareness."""
        from docx import Document as DocxDocument

        doc = DocxDocument(file_path)

        # Extract with paragraph-level structure
        paragraphs = []
        current_section = []
        section_count = 0

        for para in doc.paragraphs:
            if not para.text.strip():
                continue

            # Detect headings for structure-aware chunking
            if para.style.name.startswith("Heading"):
                # Flush current section
                if current_section:
                    section_text = "\n\n".join(current_section)
                    section_count += 1
                    paragraphs.append(PageContent(page_number=section_count, text=section_text))
                    current_section = []
                current_section.append(f"## {para.text}")
            else:
                current_section.append(para.text)

        # Flush remaining section
        if current_section:
            section_count += 1
            section_text = "\n\n".join(current_section)
            paragraphs.append(PageContent(page_number=section_count, text=section_text))

        # If no headings detected, treat as single page
        if not paragraphs:
            full_text = "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())
            paragraphs = [PageContent(page_number=1, text=full_text)]
            section_count = 1

        return ExtractedDocument(
            pages=paragraphs,
            total_pages=section_count,
            title=Path(file_path).stem,
        )

    def _extract_image_ocr(self, file_path: str) -> ExtractedDocument:
        """Extract text from images using OCR (Tesseract)."""
        try:
            import pytesseract
            from PIL import Image

            image = Image.open(file_path)
            text = _sanitize_text(pytesseract.image_to_string(image))

            if not text.strip():
                logger.warning("ocr_no_text_found", file_path=file_path)
                return ExtractedDocument(
                    pages=[],
                    total_pages=1,
                    title=Path(file_path).stem,
                )

            logger.info("ocr_extracted", file_path=file_path, chars=len(text))
            return ExtractedDocument(
                pages=[PageContent(page_number=1, text=text.strip())],
                total_pages=1,
                title=Path(file_path).stem,
            )

        except ImportError:
            logger.error("ocr_not_available", message="pytesseract or Pillow not installed")
            raise ValueError("OCR is not available. Install pytesseract and Pillow.")

    def _ocr_pdf_page(self, page) -> str:
        """OCR a single PDF page (for image-based pages)."""
        try:
            import pytesseract
            from PIL import Image
            import io

            # Render page to image
            pix = page.get_pixmap(dpi=200)
            img_data = pix.tobytes("png")
            image = Image.open(io.BytesIO(img_data))
            text = _sanitize_text(pytesseract.image_to_string(image))
            return text

        except ImportError:
            return ""
        except Exception as e:
            logger.warning("ocr_page_failed", error=str(e))
            return ""
